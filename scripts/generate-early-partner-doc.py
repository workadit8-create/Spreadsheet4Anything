#!/usr/bin/env python3
"""Generate Word doc: Early Partner Program management."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "docs")
OUT_PATH = os.path.join(OUT_DIR, "Early-Partner-Management.docx")


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(1)
        for r in p.runs:
            r.font.name = "Consolas"
            r.font.size = Pt(10)


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_heading("Early Partner Management", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Program pilot — 5 client pertama · Akuntansi App")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Versi: Juni 2025 | Internal developer + template komunikasi partner")
    doc.add_paragraph()

    add_heading(doc, "1. Tujuan program", 1)
    add_bullets(doc, [
        "5 perusahaan pertama sebagai design partner (bukan hanya user).",
        "Semua addon (inventory, pajak, CRM, proyek, aset, mobile) gratis selama periode pilot.",
        "Kumpulkan feedback nyata: addon mana dipakai, mana tidak.",
        "Validasi produk sebelum harga addon untuk client 6+.",
        "Bangun case study & referensi untuk penjualan.",
    ])

    add_heading(doc, "2. Syarat Early Partner", 1)
    add_table(doc,
        ["Kriteria", "Aturan"],
        [
            ["Kuota", "Maksimal 5 perusahaan (slot EP-01 s/d EP-05)"],
            ["Segmen", "UMKM dagang / jasa / manufaktur (≤20 karyawan)"],
            ["Komitmen partner", "Pakai app rutin, feedback bulanan, toleransi versi MVP"],
            ["Komitmen developer", "Core stabil + addon pilot bertahap, support sesuai SLA"],
            ["Bukan", "Custom development gratis tanpa batas; janji fitur enterprise"],
        ],
    )

    add_heading(doc, "3. Paket yang diberikan (gratis selama pilot)", 1)
    add_table(doc,
        ["Komponen", "Pilot", "Setelah pilot (client 6+)"],
        [
            ["Core (transaksi, laporan, posting)", "Sesuai kesepakatan core", "Harga core standar"],
            ["Addon Inventory", "Gratis", "Per modul / bundle"],
            ["Addon Pajak", "Gratis", "Per modul / bundle"],
            ["Addon CRM", "Gratis", "Per modul / bundle"],
            ["Addon Project / job costing", "Gratis", "Per modul / bundle"],
            ["Addon Aset management", "Gratis", "Per modul / bundle"],
            ["Mobile / PWA", "Gratis", "Termasuk atau addon"],
        ],
    )
    add_para(doc, "Addon = versi MVP; dirilis bertahap, tidak semua tersedia di hari pertama go-live.")

    add_heading(doc, "4. Periode & transisi ke harga normal", 1)
    add_table(doc,
        ["Item", "Rekomendasi"],
        [
            ["Durasi pilot addon gratis", "12 bulan dari tanggal go-live resmi"],
            ["Pemberitahuan sebelum berakhir", "60 hari sebelum (email/WhatsApp)"],
            ["Opsi setelah pilot", "Lanjut bayar addon / bundle diskon early partner / core saja"],
            ["Diskon loyalitas (opsional)", "10–15% addon jika lanjut dalam 30 hari setelah pilot"],
        ],
    )

    add_heading(doc, "5. Template kesepakatan (WhatsApp / email ke partner)", 1)
    agreement = (
        "Halo [Nama],\n\n"
        "Terima kasih menjadi Early Partner Akuntansi App (slot [EP-0X]).\n\n"
        "Yang Anda dapatkan:\n"
        "• Aplikasi core (transaksi, laporan, posting) sesuai paket kami\n"
        "• Semua addon pilot GRATIS selama [12] bulan dari go-live ([tanggal])\n"
        "  (Inventory, Pajak, CRM, Proyek, Aset, akses mobile — versi MVP, rilis bertahap)\n\n"
        "Yang kami harapkan:\n"
        "• Tim Anda pakai aplikasi untuk operasi nyata (bukan hanya demo)\n"
        "• Feedback singkat 1x per bulan (15 menit call/chat)\n"
        "• Melaporkan bug/kendala lewat template lapor bug\n"
        "• Memahami ini versi pilot — ada perbaikan berkala\n\n"
        "Setelah periode pilot, addon akan berbayar per modul; kami akan info "
        "60 hari sebelumnya. Harga core tetap sesuai kesepakatan awal.\n\n"
        "Konfirmasi: balas \"SETUJU\" jika Anda menerima syarat di atas.\n\n"
        "Terima kasih,\n[Nama kamu]"
    )
    p = doc.add_paragraph(agreement)
    p.paragraph_format.left_indent = Cm(1)

    add_heading(doc, "6. Tracker slot Early Partner (isi internal)", 1)
    add_para(doc, "Salin tabel ke Google Sheet atau gunakan file: docs/Early-Partner-Tracker.csv")
    add_table(doc,
        ["Slot", "Perusahaan", "Segmen", "Go-live", "Pilot berakhir", "Owner email", "Status"],
        [
            ["EP-01", "Client 1 (aktif)", "…", "…", "+12 bln", "…", "Aktif"],
            ["EP-02", "", "", "", "", "", "Kosong"],
            ["EP-03", "", "", "", "", "", "Kosong"],
            ["EP-04", "", "", "", "", "", "Kosong"],
            ["EP-05", "", "", "", "", "", "Kosong"],
        ],
    )

    add_heading(doc, "7. Checklist onboarding Early Partner", 1)
    add_numbered(doc, [
        "Tandai slot EP-0X di tracker (kosong → reserved → aktif).",
        "Kirim template kesepakatan (section 5), minta konfirmasi SETUJU.",
        "Onboarding teknis: instance terpisah (database + URL) — lihat Onboarding-Client-Baru.docx.",
        "Setup owner + users + training singkat.",
        "Kirim URL, template lapor bug, ekspektasi performa.",
        "Catat tanggal go-live → hitung pilot berakhir otomatis.",
        "Jadwalkan feedback bulanan (tanggal tetap, mis. tiap tanggal 15).",
    ])

    add_heading(doc, "8. Feedback bulanan (15 menit — template pertanyaan)", 1)
    add_table(doc,
        ["Pertanyaan", "Catatan partner"],
        [
            ["Modul apa dipakai paling sering minggu ini?", ""],
            ["Modul apa tidak dipakai / kenapa?", ""],
            ["1 kendala terbesar?", ""],
            ["Addon pilot mana paling dibutuhkan berikutnya?", ""],
            ["Skor kepuasan 1–10", ""],
        ],
    )

    add_heading(doc, "9. Addon pilot — status rilis (update manual)", 1)
    add_table(doc,
        ["Addon", "Status", "Catatan rilis"],
        [
            ["Inventory", "Belum / MVP / Stabil", ""],
            ["Pajak", "Belum / MVP / Stabil", ""],
            ["CRM", "Belum / MVP / Stabil", ""],
            ["Project", "Belum / MVP / Stabil", ""],
            ["Aset", "Belum / MVP / Stabil", ""],
            ["Mobile PWA", "Belum / MVP / Stabil", ""],
        ],
    )
    add_para(doc, "Semua EP dapat akses saat addon status MVP; tidak perlu billing terpisah selama pilot.")

    add_heading(doc, "10. Batas support Early Partner", 1)
    add_table(doc,
        ["Termasuk", "Tidak termasuk"],
        [
            ["Bug fix, deploy, smoke test", "Custom fitur di luar roadmap addon"],
            ["Edukasi penggunaan & workaround", "Integrasi pihak ketiga bespoke"],
            ["Koreksi data karena bug sistem", "Migrasi data historis 5+ tahun"],
            ["SLA mitigasi bug (P0/P1/P2)", "Training berulang tanpa batas"],
        ],
    )

    add_heading(doc, "11. Transisi setelah pilot (60 hari sebelum)", 1)
    add_numbered(doc, [
        "Kirim email transisi: tanggal berakhir, harga addon/bundle, opsi diskon loyalitas.",
        "Tawarkan call 30 menit untuk review pemakaian addon.",
        "Partner pilih: lanjut semua addon / sebagian / core saja.",
        "Update tracker: status → Lanjut / Downgrade / Churn.",
        "Nonaktifkan modul addon di SETTING jika partner tidak lanjut (saat billing aktif).",
    ])
    add_para(doc, "Template email transisi:", bold=True)
    transisi = (
        "Pilot Early Partner Anda berakhir pada [tanggal].\n"
        "Addon yang aktif: [daftar]. Harga usulan: [core] + [addon] = [total]/bulan.\n"
        "Diskon loyalitas 15% jika konfirmasi sebelum [tanggal].\n"
        "Silakan pilih paket atau diskusikan di call [tanggal/jam]."
    )
    p = doc.add_paragraph(transisi)
    p.paragraph_format.left_indent = Cm(1)

    add_heading(doc, "12. KPI program (review tiap kuartal)", 1)
    add_table(doc,
        ["KPI", "Target"],
        [
            ["Slot terisi", "5/5"],
            ["Partner aktif pakai core (min. 20 transaksi/bulan)", "≥80% partner"],
            ["Feedback bulanan terkumpul", "100% partner"],
            ["Addon dengan adoption >50% partner", "Identifikasi top 2"],
            ["Churn setelah pilot", "<20%"],
            ["Case study / referensi", "≥2 partner bersedia"],
        ],
    )

    add_heading(doc, "13. File & referensi terkait", 1)
    add_bullets(doc, [
        "docs/Early-Partner-Tracker.csv — tracker spreadsheet",
        "docs/Onboarding-Client-Baru-Akuntansi-App.docx",
        "docs/Mitigasi-Penanganan-Bugs-Akuntansi-App.docx",
        "docs/Template-Form-Lapor-Bug.docx",
        "clients/ — deploy per instance partner",
    ])

    doc.add_paragraph()
    add_para(doc, "— Regenerate: python3 scripts/generate-early-partner-doc.py")

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
