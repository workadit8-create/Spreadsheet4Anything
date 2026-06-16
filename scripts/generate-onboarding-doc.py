#!/usr/bin/env python3
"""Generate Word doc: onboarding client baru Akuntansi App."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "docs")
OUT_PATH = os.path.join(OUT_DIR, "Onboarding-Client-Baru-Akuntansi-App.docx")


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(1)
        for r in p.runs:
            r.font.name = "Consolas"
            r.font.size = Pt(10)


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Onboarding Client Baru", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Akuntansi App (Spreadsheet4Anything)")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Versi dokumen: Juni 2025 | Model: satu instance per client")
    doc.add_paragraph()

    add_heading(doc, "Prinsip penting", 1)
    add_bullets(doc, [
        "Satu client = satu database spreadsheet + satu backend engine + satu Apps Script project + satu URL web app.",
        "Client 1 dan Client 2 TIDAK berbagi database atau URL (data keuangan tidak boleh campur).",
        "Kode aplikasi bisa sama (satu repo GitHub); deploy dilakukan per project client.",
        "Update bug fix: deploy ke setiap client yang aktif.",
    ])

    add_heading(doc, "1. Bisnis & persiapan", 1)
    add_table(doc,
        ["Item", "Catatan"],
        [
            ["Kontrak / paket", "Training, support, SLA"],
            ["Email Google", "Owner + staff + akuntan (wajib akun Google)"],
            ["Nama internal", "Untuk folder clients/ dan catatan deploy"],
            ["Dokumen mitigasi bug", "docs/Mitigasi-Penanganan-Bugs-Akuntansi-App.docx"],
        ],
    )

    add_heading(doc, "2. Spreadsheet & backend (clone)", 1)
    add_numbered(doc, [
        "Buat / duplicate spreadsheet DATABASE client baru (struktur sheet sama, tanpa data transaksi client 1).",
        "Buat / duplicate BACKEND ENGINE untuk client baru (dari template BACKENDengine).",
        "Catat ID spreadsheet database dan backend (dari URL Google Sheets).",
        "Share kedua spreadsheet ke email developer (Editor).",
    ])
    add_para(doc, "Sheet yang harus ada di database:", bold=True)
    add_bullets(doc, [
        "PEMASUKAN, PEMBELIAN, MUTASI_DANA, QUOTATION, PURCHASE_REQUEST, JURNAL_MANUAL, SETTING, USERS, Master/MASTER, dll.",
    ])

    add_heading(doc, "3. Hubungkan client ↔ backend", 1)
    add_para(doc, "Di sheet SETTING (database client baru):", bold=True)
    add_table(doc,
        ["Key", "Nilai"],
        [
            ["BACKEND_ENGINE_ID", "ID spreadsheet backend client ini"],
            ["BACKEND_WEBAPP_URL", "URL web app backend client ini"],
            ["BACKEND_API_KEY", "API key backend (sama pola client 1)"],
        ],
    )
    add_para(doc, "Di konfigurasi backend engine client baru:", bold=True)
    add_bullets(doc, [
        "Set CLIENT_SPREADSHEET_ID = ID database client baru.",
    ])
    add_para(doc, "Tanpa langkah ini, laporan, posting, dan COA tidak akan sinkron.")

    add_heading(doc, "4. Clone Apps Script project (web app)", 1)
    add_numbered(doc, [
        "Buka Google Apps Script project client 1.",
        "File → Buat salinan project (atau clasp clone ke folder clients/client2).",
        "Di Config.js salinan: ganti DATABASE_ID ke ID database client baru.",
        "Deploy → New deployment → Web app.",
        "Execute as: User accessing the web app.",
        "Who has access: Anyone (atau domain jika perlu).",
        "Catat: scriptId, deployment ID (user), URL web app.",
    ])
    add_para(doc, "Di repo lokal — folder clients/client2:", bold=True)
    add_code_block(doc, [
        "cd clients/client2",
        "cp client.env.example client.env",
        "# Edit Config.js, .clasp.json, client.env",
        "chmod +x deploy.sh",
    ])

    add_heading(doc, "5. Isi konfigurasi clients/client2", 1)
    add_table(doc,
        ["File", "Isi"],
        [
            ["Config.js", "DATABASE_ID = ID spreadsheet database client"],
            [".clasp.json", "scriptId = ID Apps Script project client"],
            ["client.env", "CLIENT_NAME, DATABASE_ID, CLASP_DEPLOY_ID, BACKEND_*"],
        ],
    )
    add_para(doc, "client.env tidak di-commit ke git (rahasia deploy).")

    add_heading(doc, "6. Folder Drive upload (opsional)", 1)
    add_bullets(doc, [
        "Buat folder Drive terpisah per client untuk bukti invoice/pembelian.",
        "Update ID folder di kode atau SETTING agar file tidak campur dengan client lain.",
    ])

    add_heading(doc, "7. Deploy pertama", 1)
    add_code_block(doc, [
        "cd /Users/arthamas/Spreadsheet4Anything/clients/client2",
        "./deploy.sh \"setup client baru\"",
    ])
    add_para(doc, "Script sync kode dari repo root, push ke Apps Script client 2, redeploy URL tetap.")

    add_heading(doc, "8. Setup owner & user", 1)
    add_numbered(doc, [
        "Buka URL web app client baru (bukan URL client 1).",
        "Login dengan Google owner client.",
        "Ikuti halaman setup owner pertama.",
        "Daftarkan staff di menu Users (role: staff / akuntan).",
        "Pastikan spreadsheet di-share ke email user atau auto-share aktif.",
    ])

    add_heading(doc, "9. QA sebelum go-live", 1)
    add_table(doc,
        ["Cek", "OK?"],
        [
            ["Login owner & staff", ""],
            ["Simpan quotation / PR tanpa file", ""],
            ["Simpan invoice / PO", ""],
            ["Laporan tampil benar", ""],
            ["Posting ke backend", ""],
            ["Performa save ~2–4 detik (setelah warm-up)", ""],
            ["Smoke test / QaSmokeTest jika ada", ""],
        ],
    )

    add_heading(doc, "10. Serah terima ke client", 1)
    add_bullets(doc, [
        "URL web app client (beda dari client lain).",
        "Daftar akun + role.",
        "Singkat: refresh jika ada update; lapor bug pakai template.",
        "Dokumen mitigasi bug (opsional cetak/PDF).",
    ])
    add_para(doc, "Template email/WhatsApp:", bold=True)
    p = doc.add_paragraph(
        "Aplikasi Akuntansi sudah siap digunakan.\n"
        "URL: [link web app client ini]\n"
        "Login: gunakan akun Google yang sudah didaftarkan.\n"
        "Setelah buka app, tunggu dashboard load sebelum transaksi pertama.\n"
        "Jika ada kendala, kirim screenshot + langkah + jam kejadian."
    )
    p.paragraph_format.left_indent = Cm(1)

    add_heading(doc, "11. Maintenance multi-client", 1)
    add_table(doc,
        ["Aktivitas", "Client 1 (root)", "Client 2"],
        [
            ["Bug fix / fitur", "./scripts/deploy.sh \"...\"" , "clients/client2/deploy.sh \"...\""],
            ["Kode sumber", "Edit *.js / index.html di repo root", "Sync otomatis saat deploy"],
            ["Config client", "Config.js di root", "clients/client2/Config.js"],
            ["Executions log", "Apps Script project client 1", "Apps Script project client 2"],
            ["Backup git", "Satu repo GitHub", "Satu repo GitHub"],
        ],
    )

    add_heading(doc, "12. Buat client 3, 4, …", 1)
    add_code_block(doc, [
        "cp -R clients/_template clients/client3",
        "cd clients/client3",
        "cp client.env.example client.env",
        "# Ulangi checklist dari langkah 2",
    ])

    add_heading(doc, "Estimasi waktu", 1)
    add_table(doc,
        ["Tahap", "Durasi (template sudah ada)"],
        [
            ["Spreadsheet + backend clone", "1–2 jam"],
            ["Apps Script clone + deploy", "30–60 menit"],
            ["QA + training singkat", "2–4 jam"],
            ["Total", "~1 hari kerja"],
        ],
    )

    add_heading(doc, "Yang TIDAK dilakukan", 1)
    add_table(doc,
        ["Jangan", "Alasan"],
        [
            ["User client 2 di USERS client 1", "Data & akses campur"],
            ["Satu URL untuk banyak perusahaan", "Tidak didukung arsitektur"],
            ["Satu database untuk banyak client", "Risiko data keuangan"],
            ["Deploy root saja lalu anggap semua client update", "Project Apps Script terpisah"],
        ],
    )

    doc.add_paragraph()
    add_para(doc, "— Dokumen internal developer. Folder template: clients/_template dan clients/client2")

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
