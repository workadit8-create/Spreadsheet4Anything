#!/usr/bin/env python3
"""Generate Word doc: checklist smoke test 5 menit setelah deploy."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "docs")
OUT_PATH = os.path.join(OUT_DIR, "Checklist-Smoke-Test-5-Menit.docx")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    t = doc.add_heading("Checklist Smoke Test (5 Menit)", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Jalankan SETELAH setiap clasp push + redeploy")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Akuntansi App | Centang OK / catat jika gagal")
    doc.add_paragraph()

    doc.add_heading("Info deploy (isi sebelum tes)", 1)
    add_table(doc,
        ["Field", "Isi"],
        [
            ["Tanggal / jam", ""],
            ["Client", "Client 1 / Client 2 / …"],
            ["Versi deploy", "deskripsi redeploy"],
            ["Tester", ""],
            ["URL web app", ""],
        ],
    )

    doc.add_heading("1. Akses & sesi (~1 menit)", 1)
    add_table(doc,
        ["Cek", "OK?", "Catatan"],
        [
            ["Hard refresh browser (Ctrl/Cmd+Shift+R)", "", ""],
            ["Login Google berhasil", "", ""],
            ["Dashboard load tanpa error", "", ""],
            ["Menu sesuai role (staff/akuntan/owner)", "", ""],
        ],
    )

    doc.add_heading("2. Transaksi inti (~2 menit, tanpa upload file)", 1)
    add_table(doc,
        ["Cek", "OK?", "Catatan"],
        [
            ["Simpan Quotation (1–2 item)", "", ""],
            ["Simpan Purchase Request (1–2 item)", "", ""],
            ["Simpan Invoice baru (tunai, 1–2 item)", "", ""],
            ["Simpan PO / Pembelian (jika relevan)", "", ""],
            ["Performa save ~2–4 detik (setelah warm-up)", "", ""],
        ],
    )

    doc.add_heading("3. Integrasi backend (~1 menit)", 1)
    add_table(doc,
        ["Cek", "OK?", "Catatan"],
        [
            ["Laporan buka tanpa error", "", ""],
            ["Posting / status POST normal", "", ""],
            ["Tidak ada error di Apps Script Executions", "", ""],
        ],
    )

    doc.add_heading("4. Regresi cepat (jika bug fix spesifik)", 1)
    doc.add_paragraph(
        "Ulangi langkah persis dari laporan bug client. Jika fix untuk menu tertentu, "
        "tes hanya menu itu + 1 transaksi lain sebagai sanity check."
    )
    add_table(doc,
        ["Menu / bug yang difix", "OK?", "Catatan"],
        [
            ["", "", ""],
            ["", "", ""],
        ],
    )

    doc.add_heading("Hasil", 1)
    add_table(doc,
        ["Status", "Tindakan"],
        [
            ["Semua OK", "Kabari client boleh refresh / lanjut pakai"],
            ["Ada gagal", "Jangan kabari client — fix ulang, redeploy, tes lagi"],
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph("— Internal developer. Multi-client: jalankan per URL/client yang di-deploy.")
    p.runs[0].font.size = Pt(10)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
