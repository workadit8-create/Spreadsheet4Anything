#!/usr/bin/env python3
"""Generate Word doc: mitigasi penanganan bugs Spreadsheet4Anything."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "docs")
OUT_PATH = os.path.join(OUT_DIR, "Mitigasi-Penanganan-Bugs-Akuntansi-App.docx")


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Mitigasi & Penanganan Bugs", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Akuntansi App (Spreadsheet4Anything)")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(12)
    doc.add_paragraph("Versi dokumen: Juni 2025 | Deployment: user (URL tetap)")
    doc.add_paragraph()

    add_heading(doc, "1. Saat client lapor — kumpulkan info ini", 1)
    add_para(doc, "Minta client mengisi informasi berikut (boleh copy template WhatsApp/email di bawah):")
    add_bullets(doc, [
        "Siapa: email akun Google yang dipakai",
        "Jam kejadian (WIB)",
        "Menu / transaksi (Invoice, PO, Quotation, PR, laporan, dll.)",
        "Langkah dari buka app sampai error (step by step)",
        "Screenshot error + nomor dokumen (INV/QT/PR/PO) jika ada",
        "Ada upload file atau tidak",
        "Sekali saja atau berulang",
    ])
    add_para(doc, "Template WhatsApp/email untuk client:", bold=True)
    template = (
        "Halo, mohon info bug berikut:\n"
        "1) Email akun:\n"
        "2) Jam kejadian (WIB):\n"
        "3) Menu/transaksi:\n"
        "4) Langkah-langkah:\n"
        "5) Screenshot (jika ada):\n"
        "6) Nomor dokumen (jika ada):\n"
        "7) Upload file? (ya/tidak)\n"
        "8) Berulang? (ya/tidak)"
    )
    p = doc.add_paragraph(template)
    p.paragraph_format.left_indent = Cm(1)

    add_heading(doc, "2. Triage (prioritas — ~5 menit)", 1)
    add_table(doc,
        ["Level", "Contoh", "Respon target"],
        [
            ["P0 — Blokir", "Tidak bisa save, data hilang, angka laporan salah", "Fix hari itu; kabari client"],
            ["P1 — Ganggu", "Save sangat lambat, error kadang-kadang", "Fix 1–3 hari kerja"],
            ["P2 — Minor", "Label, layout, wishlist", "Masuk backlog / update berkala"],
        ],
    )

    add_heading(doc, "3. Reproduksi bug di sisi developer", 1)
    add_numbered(doc, [
        "Login dengan akun staff test atau role yang sama dengan client.",
        "Ikuti langkah persis seperti laporan client — tanpa upload dulu, lalu dengan upload jika relevan.",
        "Buka Google Apps Script project → Executions → cek error, stack trace, dan timestamp.",
        "Buka spreadsheet database → cek baris terkait (PEMASUKAN, PEMBELIAN, QUOTATION, dll.).",
        "Jika tidak bisa reproduksi: minta video singkat screen recording.",
    ])

    add_heading(doc, "4. Fix di kode (lokal)", 1)
    add_para(doc, "Jalankan di folder project Spreadsheet4Anything:")
    code_lines = [
        "cd /Users/arthamas/Spreadsheet4Anything",
        "# Edit file .js atau index.html sesuai bug",
        "clasp push --force",
        "clasp redeploy AKfycbzgw08PULf6FhWjiA4FlIqrRhuikcpwNnvIt02sD9I8rLzL0WprwATGTsWsdk_-TsQt --description \"fix: deskripsi singkat\"",
    ]
    for line in code_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(1)
        for r in p.runs:
            r.font.name = "Consolas"
            r.font.size = Pt(10)

    add_bullets(doc, [
        "URL deployment user tetap sama — client tidak perlu ganti link.",
        "Setelah redeploy, hard refresh browser (Ctrl/Cmd + Shift + R).",
        "Commit ke GitHub setelah fix stabil (backup).",
    ])

    add_heading(doc, "5. Deploy ke client (tanpa ganggu operasi)", 1)
    add_para(doc, "Yang client perlu tahu:", bold=True)
    add_bullets(doc, [
        "Tidak perlu install ulang atau ganti URL.",
        "Refresh halaman atau buka tab baru setelah developer konfirmasi fix.",
        "Hard refresh dulu; jika masih error, tab baru atau logout-login Google.",
    ])
    add_para(doc, "Langkah developer:", bold=True)
    add_numbered(doc, [
        "Fix kode → clasp push → redeploy deployment user.",
        "Smoke test 1–2 menit (transaksi yang error).",
        "Kabari client: sudah deploy, silakan refresh dan coba lagi.",
        "Jika bug data (baris sheet rusak): koreksi data + fix kode — jangan hanya redeploy.",
    ])

    add_heading(doc, "6. Template balas client", 1)
    add_para(doc, "Setelah fix:", bold=True)
    p = doc.add_paragraph(
        "Bug [menu/transaksi] sudah diperbaiki dan sudah di-deploy. "
        "Silakan refresh halaman (Ctrl+Shift+R) lalu coba lagi [langkah]. "
        "Jika masih error, kirim screenshot + jam kejadian."
    )
    p.paragraph_format.left_indent = Cm(1)

    add_para(doc, "Belum bisa fix cepat:", bold=True)
    p = doc.add_paragraph(
        "Sudah dicatat. Untuk sementara workaround: [mis. save tanpa file / pakai menu lain]. "
        "Estimasi perbaikan: [hari]. Data di sheet aman; yang bermasalah hanya [UI/proses X]."
    )
    p.paragraph_format.left_indent = Cm(1)

    add_heading(doc, "7. Dokumentasi internal", 1)
    add_bullets(doc, [
        "Tanggal lapor",
        "Client / user email",
        "Root cause (1 kalimat)",
        "File yang diubah",
        "Versi deploy (v193, v194, …)",
        "Status: fixed / workaround / backlog",
    ])
    add_para(doc, "Catat di GitHub issue atau notes — repo GitHub sebagai jejak historis.")

    add_heading(doc, "8. Kasus khusus", 1)
    add_table(doc,
        ["Kasus", "Penanganan"],
        [
            ["Angka laporan salah", "Audit data + posting; bisa butuh koreksi jurnal — tidak cukup redeploy saja"],
            ["Nomor dokumen double", "Cek lock, counter PropertiesService, baris duplikat di sheet"],
            ["Pesan \"Sistem sibuk\"", "Biasanya 2 orang save bersamaan; edukasi client atau review lock"],
            ["Permission / tidak bisa login", "Cek sheet USERS, share database, deploy Execute as USER_ACCESSING"],
            ["Lambat terus (>8s)", "Bukan selalu bug — ekspektasi 2–4s normal; pertimbangkan Firestore fase berikutnya"],
        ],
    )

    add_heading(doc, "9. SLA yang realistis untuk client", 1)
    add_table(doc,
        ["Level", "Respon", "Perbaikan"],
        [
            ["P0", "< 4 jam kerja", "Same day jika bisa"],
            ["P1", "1 hari kerja", "1–3 hari kerja"],
            ["P2", "Masuk backlog", "Update mingguan / bulanan"],
        ],
    )

    add_heading(doc, "10. Alur singkat (checklist)", 1)
    add_numbered(doc, [
        "Client lapor → kumpulkan info (section 1)",
        "Triage P0/P1/P2 (section 2)",
        "Reproduksi + cek Executions log (section 3)",
        "Fix kode + push + redeploy user (section 4)",
        "Smoke test (section 5)",
        "Kabari client refresh (section 6)",
        "Catat di GitHub/notes (section 7)",
    ])

    add_heading(doc, "Referensi teknis", 1)
    add_bullets(doc, [
        "Repo: github.com/workadit8-create/Spreadsheet4Anything",
        "Web app: deployment user (URL tidak berubah setelah redeploy)",
        "Auth: Execute as USER_ACCESSING",
        "Ekspektasi performa transaksi tanpa file: ~2–4 detik setelah warm-up",
    ])

    doc.add_paragraph()
    add_para(doc, "— Dokumen ini dibuat untuk internal developer. Perbarui jika proses deploy atau SLA berubah.", bold=False)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
