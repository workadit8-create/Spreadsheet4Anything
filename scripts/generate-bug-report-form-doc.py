#!/usr/bin/env python3
"""Generate Word doc: template form lapor bug (Google Form + WhatsApp)."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "docs")
OUT_PATH = os.path.join(OUT_DIR, "Template-Form-Lapor-Bug.docx")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    t = doc.add_heading("Template Form Lapor Bug", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Untuk client + panduan buat Google Form")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading("A. Template WhatsApp / email (kirim ke client)", 1)
    template = (
        "Halo, jika ada kendala di Aplikasi Akuntansi, mohon kirim info berikut:\n\n"
        "1. Email akun Google yang dipakai login:\n"
        "2. Jam kejadian (WIB):\n"
        "3. Menu / transaksi (Invoice, PO, laporan, dll.):\n"
        "4. Langkah dari buka app sampai error:\n"
        "5. Screenshot error (jika ada):\n"
        "6. Nomor dokumen (INV/QT/PR/PO) jika ada:\n"
        "7. Upload file bukti? (ya/tidak):\n"
        "8. Error sekali atau berulang? (ya/tidak):\n\n"
        "Terima kasih — tim kami akan respon sesuai tingkat urgensi."
    )
    p = doc.add_paragraph(template)
    p.paragraph_format.left_indent = Cm(1)

    doc.add_heading("B. Buat Google Form (untuk client)", 1)
    doc.add_paragraph(
        "Buka Google Form → buat form baru → judul: \"Lapor Kendala Aplikasi Akuntansi\". "
        "Tambahkan pertanyaan berikut (semua wajib kecuali screenshot):"
    )
    add_table(doc,
        ["No", "Pertanyaan", "Tipe Form", "Opsi / catatan"],
        [
            ["1", "Email akun Google", "Short answer", "Validasi: wajib"],
            ["2", "Nama perusahaan", "Short answer", ""],
            ["3", "Jam kejadian (WIB)", "Short answer", "Contoh: 14:30"],
            ["4", "Menu / transaksi", "Dropdown", "Dashboard, Invoice, PO, Quotation, PR, Laporan, Login, Lainnya"],
            ["5", "Langkah-langkah sampai error", "Paragraph", "Step by step"],
            ["6", "Nomor dokumen (jika ada)", "Short answer", "INV-/QT-/PR-/PO-"],
            ["7", "Upload file bukti transaksi?", "Multiple choice", "Ya / Tidak"],
            ["8", "Error berulang?", "Multiple choice", "Ya / Tidak / Kadang-kadang"],
            ["9", "Screenshot error", "File upload", "Opsional"],
            ["10", "Tingkat urgensi (client)", "Multiple choice", "Tidak bisa input / Data salah / Lambat / Tampilan saja"],
        ],
    )
    doc.add_paragraph("Setelah publish: bagikan link form ke client (bisa di footer app atau onboarding).")

    doc.add_heading("C. Notifikasi ke developer", 1)
    for item in [
        "Google Form → Settings → Get email notifications for new responses.",
        "Atau: Responses → Link to Sheets → pantau sheet otomatis.",
        "Triase dengan dokumen: Mitigasi-Penanganan-Bugs-Akuntansi-App.docx",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("D. Mapping urgensi → SLA (internal)", 1)
    add_table(doc,
        ["Client pilih", "Level internal", "Target respon"],
        [
            ["Tidak bisa input / data hilang", "P0", "Same day"],
            ["Angka laporan salah", "P0", "Same day + audit data"],
            ["Lambat / kadang error", "P1", "1–3 hari"],
            ["Tampilan / label", "P2", "Backlog"],
        ],
    )

    doc.add_paragraph()
    doc.add_paragraph("— Lampirkan link form ini saat serah terima client baru.")

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
